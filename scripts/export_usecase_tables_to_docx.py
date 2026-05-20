from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"D:\Code\TRAE\bishe")
SOURCE = ROOT / "note" / "draft3" / "usecase_spec_tables.md"
TARGET = ROOT / "note" / "draft3" / "用例规约表-便于复制.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=40, start=70, bottom=40, end=70) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "7F7F7F")


def style_run(run, *, size=10.5, bold=False, color="000000", font="宋体") -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key in ("ascii", "hAnsi", "eastAsia"):
        rfonts.set(qn(f"w:{key}"), font)


def add_cell_text(cell, text: str, *, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    parts = text.split("<br>")
    for idx, part in enumerate(parts):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(part.strip())
        style_run(run, size=10.5, bold=bold, color="000000", font="宋体")


def parse_tables(md: str):
    lines = md.splitlines()
    tables = []
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith("## "):
            title = line[3:].strip()
            i += 1
            while i < len(lines) and not lines[i].strip():
                i += 1
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i].rstrip())
                i += 1
            rows = []
            for raw in block:
                if re.match(r"^\|\s*-+\s*\|\s*-+\s*\|$", raw.strip()):
                    continue
                cols = [c.strip() for c in raw.strip().strip("|").split("|")]
                if len(cols) >= 2:
                    rows.append(cols[:2])
            if rows:
                tables.append((title, rows))
            continue
        i += 1
    return tables


def build_doc(tables) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("第三章用例规约表")
    style_run(run, size=14, bold=True, font="黑体")

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(6)
    intro.paragraph_format.line_spacing = 1.0
    r = intro.add_run("以下表格由 Markdown 规约内容整理生成，便于在 WPS 中直接复制到论文正文。")
    style_run(r, size=10.5, font="宋体")

    for idx, (table_title, rows) in enumerate(tables):
        if idx > 0:
            doc.add_paragraph()

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(table_title)
        style_run(run, size=10.5, bold=True, font="黑体")

        table = doc.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_table_borders(table)

        for row_idx, row_data in enumerate(rows):
            row = table.add_row()
            row.height = Cm(0.78 if row_idx == 0 else 0.96)
            row.cells[0].width = Cm(3.2)
            row.cells[1].width = Cm(12.8)

            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)

            left, right = row.cells
            if row_idx == 0:
                set_cell_shading(left, "5B9BD5")
                set_cell_shading(right, "5B9BD5")
                add_cell_text(left, row_data[0], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
                add_cell_text(right, row_data[1], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                fill = "D9E2F3" if row_idx % 2 == 1 else "EDF2F7"
                set_cell_shading(left, fill)
                set_cell_shading(right, fill)
                add_cell_text(left, row_data[0], bold=False)
                add_cell_text(right, row_data[1], bold=False)

    return doc


def main() -> None:
    md_text = SOURCE.read_text(encoding="utf-8")
    tables = parse_tables(md_text)
    if not tables:
        raise SystemExit("No markdown tables found.")
    doc = build_doc(tables)
    doc.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
